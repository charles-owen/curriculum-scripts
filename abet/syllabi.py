"""
Usage:
    syllabi [-h|--help] [--program=<program>] <syllabi_xml>

Options:
  -h --help             Show this screen.
  --program=<program>   Choose program to generator for (cs, ce, ee, me) [default: cs]
"""

#
# Script to convert syllabi from XML to Word documents
#

import os
import xml.dom.minidom

from docopt import docopt

from docx.shared import Inches
from docx.shared import Pt

from lib.syllabi.settings import Settings
from lib.syllabi.studentoutcome import StudentOutcomes
from lib.syllabi.course import Course
from lib.word import WordDocument


def process(settings: Settings):
    doc = xml.dom.minidom.parse(settings.source)

    #
    # Single file per course
    #
    for node in doc.firstChild.childNodes:
        if node.nodeType == node.ELEMENT_NODE and node.nodeName == "course":
            word = createWordDocument()
            course = process_course(settings, word, node)
            if course is not None:
                word.save(settings.program + '/' + course.short_number + ".docx")

    #
    # Single file for everything
    #
    word = createWordDocument()
    for node in doc.firstChild.childNodes:
        if node.nodeType == node.ELEMENT_NODE and node.nodeName == "course":
            course = process_course(settings, word, node)
            if course is not None:
                word.page_break()

    word.save(settings.program + "/syllabi.docx")


def createWordDocument():
    word = WordDocument()
    style = word.add_style("Syllabus")
    word.indent(style, Inches(0.25), Inches(-0.25))

    style = word.add_style("SyllabusTitle")
    word.indent(style, Inches(0.25), Inches(0))

    style = word.add_style("Syllabus1")
    word.indent(style, Inches(0.75), Inches(-0.25))
    word.before_after(style, Pt(0), Pt(0))

    style = word.add_style("Syllabus2")
    word.indent(style, Inches(1.25), Inches(-0.25))
    word.before_after(style, Pt(0), Pt(0))

    style = word.add_style("Syllabus3")
    word.indent(style, Inches(1.75), Inches(-0.25))
    word.before_after(style, Pt(0), Pt(0))

    return word


def process_course(settings: Settings, word, node):
    course = Course(settings, node)
    if not course.is_active():
        return None

    print(course.short_number)

  #  word.add_item(1, course.number + " " + course.name, "SyllabusTitle")
    word.word.add_paragraph(course.number + " " + course.name, style="SyllabusTitle")

    item2 = f"{course.credits} credit hour/{course.contact} contact hour"
    if course.type != '':
        item2 += '/' + course.type
    word.add_item(2, item2, "Syllabus")
    word.add_item(3, course.instructor, "Syllabus")

    word.add_item(4, course.text, "Syllabus")
    word.add_list(course.supplemental, "Syllabus1")

    req = ''
    if course.required:
        req = 'Required'
    elif course.elective:
        req = 'Elective'
    elif course.selected_elective:
        req = 'Selected Elective'
    elif course.math_science:
        req = 'Math and Science'

    word.add_item(5, "Specific course information", "Syllabus")
    word.add_item('a', course.description, "Syllabus1")
    word.add_item('b', course.prereq, "Syllabus1")
    word.add_item('c', req, "Syllabus1")

    add_outcomes(settings, word, 6, course.outcomes, course.studentOutcomes)

    word.add_item(7, "Topics", "Syllabus")
    word.add_list(course.topics, "Syllabus1")

    if course.topics2:
        word.word.add_paragraph(f"\tSpecific Secure Computing Topics", style="Syllabus")
        word.add_list(course.topics2, "Syllabus1")

    return course


def add_outcomes(settings: Settings, word, num, outcomes, mapping):
    word.add_item(num, "Specific goals for the course", "Syllabus")
    word.add_item('a', "Course Outcomes", "Syllabus1")

    num = 1
    for outcome in outcomes:
        roman = intToRoman(num).lower()
        word.add_item(roman, outcome.text, "Syllabus2")
        num1 = 1
        for sub in outcome.subs:
            word.add_item(num1, sub, "Syllabus3")
            num1 += 1
        num += 1

        # outcome_mapping = outcome.get_mapping(settings.program)
        # for item in outcome_mapping:
        #     mapping[item] = True

    word.add_item('b', "Student Outcomes", "Syllabus1")

    studentOutcomes = StudentOutcomes()
    num = 1
    for outcome in studentOutcomes.get(settings.program, mapping):
        word.add_paragraph(outcome, "Syllabus2")
        num += 1


# Function to calculate Roman values
def intToRoman(num):
    # Storing roman values of digits from 0-9
    # when placed at different places
    m = ["", "M", "MM", "MMM"]
    c = ["", "C", "CC", "CCC", "CD", "D",
         "DC", "DCC", "DCCC", "CM "]
    x = ["", "X", "XX", "XXX", "XL", "L",
         "LX", "LXX", "LXXX", "XC"]
    i = ["", "I", "II", "III", "IV", "V",
         "VI", "VII", "VIII", "IX"]

    # Converting to roman
    thousands = m[num // 1000]
    hundreds = c[(num % 1000) // 100]
    tens = x[(num % 100) // 10]
    ones = i[num % 10]

    ans = (thousands + hundreds +
           tens + ones)

    return ans


#
# Program entry point
#
if __name__ == '__main__':
    args = docopt(__doc__)
    #print(args)

    settings = Settings(os.getcwd(), args)
    process(settings)

