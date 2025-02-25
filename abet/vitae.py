"""
Usage:
    vitae [-h|--help] [--program=<program>] <faculty_xml>

Options:
  -h --help             Show this screen.
"""

#
# Script to convert faculty vitae from XML to Word documents
#

import os
import xml.dom.minidom

from docopt import docopt

from docx.shared import Inches
from docx.shared import Pt

from lib.vitae.settings import Settings
from lib.vitae.faculty import Faculty
from lib.word import WordDocument


def process(settings: Settings):
    doc = xml.dom.minidom.parse(settings.source)

    #
    # Single file per course
    #
    for node in doc.firstChild.childNodes:
        if node.nodeType == node.ELEMENT_NODE and node.nodeName == "faculty":
            word = createWordDocument()
            faculty = process_faculty(settings, word, node)
            if faculty is not None:
                word.save('vitae/' + faculty.user + ".docx")

    #
    # Single file for everything
    #
    word = createWordDocument()
    for node in doc.firstChild.childNodes:
        if node.nodeType == node.ELEMENT_NODE and node.nodeName == "faculty":
            course = process_faculty(settings, word, node)
            if course is not None:
                word.page_break()

    word.save("vitae/vitae.docx")


def createWordDocument():
    word = WordDocument()
    style = word.add_style("Vita")
    word.indent(style, Inches(0.25), Inches(-0.25))

    style = word.add_style("Vita1")
    word.indent(style, Inches(0.25), Inches(0))
    word.before_after(style, Pt(0), Pt(0))

    style = word.add_style("Vita2")
    word.indent(style, Inches(0.50), Inches(0))
    word.before_after(style, Pt(0), Pt(0))

    # style = word.add_style("Syllabus3")
    # word.indent(style, Inches(1.75), Inches(-0.25))
    # word.before_after(style, Pt(0), Pt(0))

    return word


def process_faculty(settings: Settings, word, node):
    faculty = Faculty(settings, node)

    print(faculty.name)

    word.add_item(1, faculty.name, "Vita")
    word.add_item(2, "Education", "Vita");
    for education in faculty.education:
        word.add_paragraph(education, "Vita1")

    word.add_item(3, "Academic Experience", "Vita");
    for academic in faculty.academic:
        word.add_paragraph(academic, "Vita1")

    word.add_item(4, "Non-academic Experience", "Vita");
    if faculty.non_academic:
        for academic in faculty.non_academic:
            word.add_paragraph(academic, "Vita1")
    else:
        word.add_paragraph("None", "Vita1")

    word.add_item(5, "Certifications or professional registrations", "Vita");
    for item in faculty.certifications:
        word.add_paragraph(item, "Vita1")

    word.add_item(6, "Current membership in professional organizations ", "Vita");
    for item in faculty.members:
        word.add_paragraph(item, "Vita1")

    word.add_item(7, "Honors and awards", "Vita");
    for item in faculty.honors:
        word.add_paragraph(item, "Vita1")

    word.add_item(8, "Service activities", "Vita");
    for item in faculty.service:
        word.add_paragraph(item, "Vita1")
    for category in faculty.categories:
        word.add_paragraph(category["title"], "Vita1")
        for service in category["service"]:
            word.add_paragraph(service, "Vita2")

    word.add_item(9, "Important publications in the past five years", "Vita");
    for item in faculty.publications:
        word.add_paragraph(item, "Vita1")

    word.add_item(10, "Recent development activities", "Vita")
    for item in faculty.development:
        word.add_paragraph(item, "Vita1")

    return faculty

#
# Program entry point
#
if __name__ == '__main__':
    args = docopt(__doc__)
    #print(args)

    settings = Settings(os.getcwd(), args)
    process(settings)

